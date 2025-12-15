from docx import Document
from docx.shared import Pt

# Neues Dokument erstellen
doc = Document()

# Titel
doc.add_heading('Unterschriftenseite Studienprotokoll', 0)
doc.add_paragraph("Dieses Dokument ist Bestandteil des Ethikantrags und wird separat eingereicht.")

# Platzhalter für Unterschriften
signatures = [
    "Prof. Dr. med. Ina Conrad-Hengerer",
    "Prof. Dr. rer. nat. Alexandra Reichenbach",
    "Prof. Dr.-Ing. Gerrit Meixner",
    "Dr. med. Lucy Kessler",
    "Jonathan Pareja Carrillo",
]

for sig in signatures:
    doc.add_paragraph(f"{sig}:\n\n\n______________________________", style="Normal")

# Hinweistext
p = doc.add_paragraph("\nHinweis: Dieses Dokument ist vom Antragsteller und den beteiligten Personen eigenhändig zu unterschreiben.")
p.runs[0].font.size = Pt(10)

# Dokument speichern
output_path = "./Unterschriftenseite_Studienprotokoll.docx"
doc.save(output_path)