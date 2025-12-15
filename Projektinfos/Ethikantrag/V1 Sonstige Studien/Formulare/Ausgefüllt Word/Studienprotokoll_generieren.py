from docx import Document

# Neues Dokument erstellen
doc = Document()

# Titel und Metadaten
doc.add_heading('Studienprotokoll / Projektplan für Ethikantrag', 0)
doc.add_paragraph("Version: 1.0 | Datum: 05.09.2025")

# 1 Projekttitel
doc.add_heading('1. Projekttitel', level=1)
doc.add_paragraph("Vergleich der klassischen kinetischen Perimetrie nach Goldmann mit einer VR-basierten Implementierung unter Nutzung von Eye-Tracking-Technologie.")

# 2 Verantwortlichkeiten
doc.add_heading('2. Verantwortlichkeiten', level=1)
doc.add_paragraph("Studienleiter/in: [Name einfügen]\nBeteiligte Wissenschaftler/innen: [Namen einfügen]\nBeteiligte Einrichtungen: Universitätsklinikum Heidelberg, Augenklinik\nRegistrierung: Gemäß Artikel 35 Deklaration von Helsinki vorgesehen.")

# 3 Wissenschaftlicher Hintergrund
doc.add_heading('3. Wissenschaftlicher Hintergrund', level=1)
doc.add_paragraph(
    "In der Augenheilkunde zählt die Perimetrie zu den Standarduntersuchungen und wird bei zahlreichen Erkrankungen eingesetzt. "
    "Die kinetische Perimetrie nach Hans Goldmann gilt seit 1945 als Goldstandard. Trotz technischer Weiterentwicklungen existieren bislang kaum validierte Alternativen, "
    "die Vorteile in Praktikabilität, Dokumentation und Verfügbarkeit bieten. Dieses Projekt untersucht, ob eine VR-Implementierung mit Eye-Tracking die gleiche diagnostische Genauigkeit wie die klassische Goldmann-Perimetrie aufweist. "
    "Dabei soll auch die Patient:innenpräferenz erfasst werden."
)

# 4 Projektziele
doc.add_heading('4. Projektziele', level=1)
doc.add_paragraph(
    "Primäres Ziel: Untersuchung, ob die Ergebnisse der VR-basierten kinetischen Perimetrie statistisch signifikant mit den Ergebnissen der Goldmann-Perimetrie übereinstimmen.\n"
    "Sekundäre Ziele: (1) Erfassung der Patient:innenpräferenz durch Fragebogen, (2) Evaluation der praktischen Durchführbarkeit des VR-Verfahrens."
)

# 5 Endpunkte
doc.add_heading('5. Endpunkte / Zielgrößen', level=1)
doc.add_paragraph(
    "Primärer Endpunkt: Abweichung der Gesichtsfeldmessung zwischen Goldmann-Perimetrie und VR-Implementierung.\n"
    "Sekundäre Endpunkte: (1) Präferenz der Patient:innen (Fragebogen), (2) Zeitbedarf je Untersuchung, (3) technische Machbarkeit und Datenqualität."
)

# 6 Studienpopulation
doc.add_heading('6. Studienpopulation', level=1)
doc.add_paragraph(
    "Gesunde Proband:innen > 18 Jahre (Phase 1, n < 10), anschließend Patient:innen > 18 Jahre mit bekannten Gesichtsfeldausfällen (Phase 2, n < 10). "
    "Der Test an gesunden Personen ermöglicht eine Baseline-Validierung und stellt sicher, dass das Verfahren technisch zuverlässig funktioniert, bevor Patient:innen mit Pathologien eingeschlossen werden."
)

# 7 Methodik und Durchführung
doc.add_heading('7. Methodik und Durchführung', level=1)
doc.add_paragraph(
    "Monozentrische Studie am Universitätsklinikum Heidelberg. \n"
    "Studiendesign: Prospektiv, explorativ. \n"
    "Ablauf: Zunächst Durchführung der VR-basierten Perimetrie an gesunden Proband:innen, anschließend an Patient:innen mit Gesichtsfeldausfällen. Alle Teilnehmer:innen durchlaufen zusätzlich eine Goldmann-Perimetrie als Vergleich. \n"
    "Datenerhebung: VR-Ergebnisse (Eye-Tracking, Eingaben), klassische Goldmann-Daten, Fragebogen. Speicherung der Daten nach DICOM-Standard."
)

# 8 Strahlenanwendung
doc.add_heading('8. Strahlenanwendung', level=1)
doc.add_paragraph("Es findet keine Strahlenanwendung statt.")

# 9 Nutzen-Risiko-Abwägung
doc.add_heading('9. Nutzen-Risiko-Abwägung', level=1)
doc.add_paragraph(
    "Die Studie ist risikoarm. Risiken sind eine mögliche Ermüdung der Augen und leichte Unannehmlichkeiten durch die VR-Brille. "
    "Der Nutzen liegt in der möglichen Etablierung eines praktikableren, validierten Verfahrens für die klinische Diagnostik."
)

# 10 Biometrie
doc.add_heading('10. Biometrie', level=1)
doc.add_paragraph(
    "Explorativer Ansatz. Die Fallzahl ist klein (n < 20), da es sich um eine Pilotstudie handelt. Statistische Auswertung der Übereinstimmung zwischen den Verfahren mittels geeigneter Vergleichstests."
)

# 11 Datenmanagement und Datenschutz
doc.add_heading('11. Datenmanagement und Datenschutz', level=1)
doc.add_paragraph(
    "Alle Daten werden pseudonymisiert gespeichert. Speicherung erfolgt lokal am Universitätsklinikum Heidelberg nach DSGVO-Standards. "
    "Zugriff haben nur autorisierte Projektmitarbeiter:innen. Datenlöschung nach Abschluss der Studie und Publikation."
)

# 12 Zeitplan
doc.add_heading('12. Zeitplan', level=1)
doc.add_paragraph(
    "Arbeitspaket 1: bis 19. Februar – Grundlagen, Auswahl Brille, Kriterien, Ethikantrag.\n"
    "Arbeitspaket 2: bis 31. Oktober – Prototypischer Algorithmus in Python.\n"
    "Arbeitspaket 3: bis 28. November – Umsetzung in VR, erste Tests an Gesunden.\n"
    "Arbeitspaket 4: 01.–12. Dezember – Klinische Tests an Patient:innen.\n"
    "Arbeitspaket 5: 15.–26. Dezember – Statistische Analyse.\n"
    "Arbeitspaket 6: bis 02. Januar – Erstellung Paper."
)

# 13 Finanzierung
doc.add_heading('13. Finanzierung und Förderung', level=1)
doc.add_paragraph("[Noch einzufügen]")

# 14 Publikation
doc.add_heading('14. Publikations- und Disseminationsplan', level=1)
doc.add_paragraph(
    "Geplante Veröffentlichung nach Abschluss der Auswertung und Benotung. Ziel: Publikation in einem peer-reviewed Journal im Bereich Ophthalmologie."
)

# 15 Weitere Anmerkungen
doc.add_heading('15. Weitere Anmerkungen', level=1)
doc.add_paragraph("Entwicklung eines standardisierten Patient:innenfragebogens zur Bewertung der Untersuchungsmethoden.")

# Dokument speichern
doc.save('./Studienprotokoll_Ethikantrag_ausgefüllt.docx')
